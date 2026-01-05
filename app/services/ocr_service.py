import io
import os
import platform
import time
import uuid
from typing import List

import pytesseract
import pypdfium2 as pdfium
from PIL import Image, ImageOps, ImageFilter
from docx import Document

from app.models.schemas import OCRResponse, PageText
from app.services import file_service
from app.core.config import settings


# ---------------------------------------------------------
# Tesseract configuration (Windows-safe)
# ---------------------------------------------------------

def configure_tesseract():
    """
    1) If env var TESSERACT_CMD is set, use it
    2) Else on Windows, fallback to default install path
    3) Else rely on PATH (Linux/macOS/Docker)
    """
    env_cmd = os.getenv("TESSERACT_CMD")
    if env_cmd and os.path.exists(env_cmd):
        pytesseract.pytesseract.tesseract_cmd = env_cmd
        return

    if platform.system().lower().startswith("win"):
        win_default = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(win_default):
            pytesseract.pytesseract.tesseract_cmd = win_default
            return


configure_tesseract()


def ocr_image(image: Image.Image) -> str:
    """
    OCR for a single image using Tesseract, with simple preprocessing.
    """
    image = image.convert("L")

    w, h = image.size
    max_side = max(w, h)
    if max_side < 1000:
        scale = 1000 / max_side
        image = image.resize((int(w * scale), int(h * scale)), Image.LANCZOS)

    image = ImageOps.autocontrast(image)
    image = image.filter(ImageFilter.SHARPEN)

    text = pytesseract.image_to_string(image)
    return text.strip()


def extract_from_pdf(file_bytes: bytes) -> List[PageText]:
    """
    Hybrid: try PDF text layer, fallback to OCR per page if empty.
    """
    pdf = pdfium.PdfDocument(file_bytes)
    pages: List[PageText] = []

    for i in range(len(pdf)):
        page = pdf[i]

        text = ""
        try:
            textpage = page.get_textpage()
            text = (textpage.get_text_range() or "").strip()
        except Exception:
            text = ""

        if not text:
            # 300 DPI render
            scale = 300 / 72.0
            bitmap = page.render(scale=scale)
            pil_image = bitmap.to_pil()
            text = ocr_image(pil_image)

        pages.append(PageText(page_number=i + 1, text=text))

    return pages


def extract_from_image(file_bytes: bytes) -> List[PageText]:
    image = Image.open(io.BytesIO(file_bytes))
    text = ocr_image(image)
    return [PageText(page_number=1, text=text)]


def extract_from_docx(file_bytes: bytes) -> List[PageText]:
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return [PageText(page_number=1, text=text)]


def process_file(
    file_bytes: bytes,
    filename: str,
    document_type: str,
    *,
    zero_retention: bool | None = None,
) -> OCRResponse:
    """
    Main OCR function.
    IMPORTANT:
    - If zero_retention=True => DO NOT store the file anywhere.
    - If zero_retention=False => store only ONE copy by filename (overwrite).
    """
    start = time.time()
    job_id = str(uuid.uuid4())

    if zero_retention is None:
        zero_retention = settings.ZERO_RETENTION_DEFAULT

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext == "pdf":
        pages = extract_from_pdf(file_bytes)
    elif ext in {"jpg", "jpeg", "png", "bmp", "tif", "tiff"}:
        pages = extract_from_image(file_bytes)
    elif ext == "docx":
        pages = extract_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext or 'unknown'}")

    full_text = "\n\n".join(p.text for p in pages)
    processing_time_ms = int((time.time() - start) * 1000)

    # ✅ FIX #1: Only store when retention is OFF
    if not zero_retention:
        file_service.save_unique_by_name(filename, file_bytes)
    else:
        # extra safety: if something old exists with same filename, remove it
        # (optional, but matches your "no retention" expectation)
        file_service.delete_if_exists(filename)

    return OCRResponse(
        job_id=job_id,
        status="success",
        document_type=document_type,
        pages=pages,
        full_text=full_text,
        metadata={
            "file_name": filename,
            "file_type": ext,
            "num_pages": len(pages),
            "processing_time_ms": processing_time_ms,
            "engine": "tesseract",
            "zero_retention": bool(zero_retention),
        },
    )
